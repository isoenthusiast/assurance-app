"""
Convert Word (.docx) and PDF files to Markdown.

Usage:
  python convert_to_md.py <input_file> [--output <output_file>]

If --output is omitted, writes result to stdout.

Supports:
  - .docx  → python-docx (preserves headings, bold, italic, tables)
  - .pdf   → PyMuPDF (fitz) (extracts text with layout preservation)
"""

import sys
import os
import argparse
import re


# ── Word (.docx) → Markdown ──────────────────────────────────────────

def docx_to_markdown(filepath: str) -> str:
    """Convert a .docx file to Markdown, preserving structure."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        text = ""
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                t = f"***{t}***"
            elif run.bold:
                t = f"**{t}**"
            elif run.italic:
                t = f"*{t}*"
            text += t

        if not text.strip():
            lines.append("")
            continue

        style = para.style.name if para.style else ""
        level = _heading_level(style, para)

        if level:
            lines.append(f"{'#' * level} {text.strip()}")
        else:
            lines.append(text)

    # Tables
    for table in doc.tables:
        lines.append("")
        rows = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        # Header row
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            # Pad with empty cells if needed
            padded = row + [""] * (len(rows[0]) - len(row))
            lines.append("| " + " | ".join(padded[:len(rows[0])]) + " |")
        lines.append("")

    return "\n".join(lines)


def _heading_level(style_name: str, para) -> int | None:
    """Determine heading level from paragraph style."""
    # Check style name
    if style_name:
        m = re.match(r'^Heading\s*(\d)', style_name, re.IGNORECASE)
        if m:
            return int(m.group(1))

    # Check outline level via XML
    try:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            outline_lvl = pPr.find(qn('w:outlineLvl'))
            if outline_lvl is not None:
                val = outline_lvl.get(qn('w:val'))
                if val is not None:
                    return int(val) + 1
    except Exception:
        pass

    return None


# ── PDF → Markdown ───────────────────────────────────────────────────

def pdf_to_markdown(filepath: str) -> str:
    """Convert a .pdf file to Markdown, extracting text page by page."""
    import fitz  # PyMuPDF

    doc = fitz.open(filepath)
    lines = []

    for page_num, page in enumerate(doc, 1):
        lines.append(f"\n## Page {page_num}\n")

        # Try blocks first (preserves layout order)
        blocks = page.get_text("blocks")
        if blocks:
            for block in blocks:
                text = block[4].strip()
                if not text:
                    lines.append("")
                    continue
                # Detect large/bold text as headings (heuristic)
                block_type = block[6] if len(block) > 6 else 0
                if block_type == 1:  # Image block
                    lines.append(f"> *[Image on page {page_num}]*")
                else:
                    lines.append(text)
        else:
            # Fallback: plain text
            text = page.get_text("text")
            lines.append(text)

    doc.close()
    return "\n".join(lines)


# ── Main ─────────────────────────────────────────────────────────────

SUPPORTED = {
    ".docx": docx_to_markdown,
    ".pdf": pdf_to_markdown,
}


def main():
    parser = argparse.ArgumentParser(description="Convert Word/PDF to Markdown")
    parser.add_argument("input", help="Input file path (.docx or .pdf)")
    parser.add_argument("--output", "-o", help="Output file path (default: stdout)")
    args = parser.parse_args()

    ext = os.path.splitext(args.input)[1].lower()
    if ext not in SUPPORTED:
        print(f"ERROR: Unsupported format '{ext}'. Supported: {', '.join(SUPPORTED)}", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(args.input):
        print(f"ERROR: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    converter = SUPPORTED[ext]
    md_content = converter(args.input)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Converted → {args.output}", file=sys.stderr)
    else:
        print(md_content)


if __name__ == "__main__":
    main()
